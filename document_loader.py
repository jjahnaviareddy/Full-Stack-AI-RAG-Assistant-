"""
Document loading utilities for PDF, TXT, DOCX, CSV files.
"""
import csv
import io
from pathlib import Path
from typing import List

from langchain.schema import Document
from loguru import logger


def load_document(file_path: str, filename: str) -> List[Document]:
    """
    Load a document from disk and return LangChain Document objects.
    Supports: PDF, TXT, DOCX, CSV, MD
    """
    path = Path(file_path)
    ext = path.suffix.lower()

    loaders = {
        ".pdf": _load_pdf,
        ".txt": _load_txt,
        ".md": _load_txt,
        ".docx": _load_docx,
        ".csv": _load_csv,
    }

    loader_fn = loaders.get(ext)
    if loader_fn is None:
        raise ValueError(f"Unsupported file type: {ext}")

    try:
        docs = loader_fn(file_path, filename)
        logger.info(f"Loaded {len(docs)} pages/sections from '{filename}'")
        return docs
    except Exception as e:
        logger.error(f"Failed to load document '{filename}': {e}")
        raise


def _load_pdf(file_path: str, filename: str) -> List[Document]:
    from pypdf import PdfReader

    reader = PdfReader(file_path)
    docs = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        if text.strip():
            docs.append(
                Document(
                    page_content=text,
                    metadata={
                        "source": filename,
                        "page": i + 1,
                        "total_pages": len(reader.pages),
                    },
                )
            )
    return docs


def _load_txt(file_path: str, filename: str) -> List[Document]:
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        text = f.read()
    return [Document(page_content=text, metadata={"source": filename})]


def _load_docx(file_path: str, filename: str) -> List[Document]:
    from docx import Document as DocxDocument

    doc = DocxDocument(file_path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    text = "\n\n".join(paragraphs)
    return [Document(page_content=text, metadata={"source": filename})]


def _load_csv(file_path: str, filename: str) -> List[Document]:
    docs = []
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        reader = csv.DictReader(f)
        for i, row in enumerate(reader):
            text = "\n".join(f"{k}: {v}" for k, v in row.items() if v)
            if text.strip():
                docs.append(
                    Document(
                        page_content=text,
                        metadata={"source": filename, "row": i + 1},
                    )
                )
    return docs


SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".md", ".docx", ".csv"}


def is_supported(filename: str) -> bool:
    return Path(filename).suffix.lower() in SUPPORTED_EXTENSIONS
